# -*- coding: utf-8 -*-
"""
09_artigo_docx.py - Ensaio analítico.
Acidentes de trabalho notificados entre profissionais da saúde em Campos dos Goytacazes (RJ), 2018-2025.
"""
import os, re, subprocess, shutil
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

RAIZ = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
os.chdir(RAIZ)

d = Document()
for s in d.sections:
    s.page_width, s.page_height = Cm(21), Cm(29.7)
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2.5)
n = d.styles["Normal"]
n.font.name = "Times New Roman"; n.font.size = Pt(12)
n.paragraph_format.space_after = Pt(0)
n.paragraph_format.line_spacing = 1.5

def _runs(p, texto, size):
    for parte in re.split(r"(\[\[i\]\].*?\[\[/i\]\]|\[\[b\]\].*?\[\[/b\]\])", texto):
        if not parte: continue
        it = parte.startswith("[[i]]"); ng = parte.startswith("[[b]]")
        limpo = re.sub(r"\[\[/?[ib]\]\]", "", parte)
        r = p.add_run(limpo); r.font.size = Pt(size); r.italic = it; r.bold = ng
    return p

def par(texto, indent=True, just=True, size=12, before=0, after=0, center=False):
    p = d.add_paragraph(); _runs(p, texto, size)
    pf = p.paragraph_format
    pf.first_line_indent = Cm(1.25) if indent else Cm(0)
    pf.space_before, pf.space_after = Pt(before), Pt(after)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else (
        WD_ALIGN_PARAGRAPH.JUSTIFY if just else WD_ALIGN_PARAGRAPH.LEFT)
    return p

def tabela(dados_tab, fonte_txt):
    t = d.add_table(rows=0, cols=len(dados_tab[0])); t.style = "Table Grid"
    for i, row in enumerate(dados_tab):
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = ""
            _runs(cells[j].paragraphs[0], v, 7.5)
            for r_ in cells[j].paragraphs[0].runs:
                r_.bold = r_.bold or (i == 0)
    par(fonte_txt, indent=False, size=8, after=4)

def figura(path, caption):
    pf = d.add_paragraph(); pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.paragraph_format.space_before = Pt(3)
    pf.add_run().add_picture(path, width=Cm(10.5))
    par(caption, indent=False, size=8, after=3)

# ============================ TÍTULO ===========================================
par("[[b]]Panorama da saúde do trabalhador em Campos dos Goytacazes (RJ)[[/b]]",
    indent=False, center=True, size=14, after=8)

# ============================ ENSAIO ===========================================
CORPO = [

 # =======================================================================
 # 1. INTRODUÇÃO
 # =======================================================================
 "Campos dos Goytacazes, maior município do estado do Rio de Janeiro, "
 "contava com 483.540 habitantes no Censo de 2022 (IBGE). Sua economia "
 "organiza-se em três ciclos. O ciclo açucareiro entrou em colapso nos anos "
 "1980. O ciclo petrolífero (1970-2014), impulsionado pela Bacia de Campos, "
 "gerou R$ 339,7 milhões em [[i]]royalties[[/i]] entre 2018 e 2025 (ANP). "
 "O terceiro ciclo, em curso desde 2014, é de estagnação. Silva e Hasenclever "
 "(2019) demonstram que a riqueza petrolífera não superou o "
 "subdesenvolvimento histórico: em 2024, 71% das receitas provieram de "
 "transferências (Siconfi/STN), com PIB per capita de R$ 88.831 contrastando "
 "com IDHM de 0,716 e 37,7% da população com até meio salário mínimo. "
 "O setor saúde responde por 15.002 postos de trabalho (CEMPRE 2024).",

 # =======================================================================
 # 2. REGIMES PREVIDENCIÁRIOS E A CAT
 # =======================================================================
 "O funcionalismo público de Campos opera sob dois regimes previdenciários. "
 "O Regime Próprio (RPPS) cobre estatutários, com contribuições patronais "
 "de R$ 74,7 milhões em 2025. O INSS cobre celetistas, com R$ 4,8 milhões "
 "no mesmo ano (Portal da Transparência). A razão RPPS/INSS passou de 3,3 "
 "para 15,5 entre 2024 e 2025. A Comunicação de Acidente de Trabalho (CAT), "
 "instituída pela Lei nº 8.213/1991, é instrumento exclusivo do INSS: "
 "acidentes de estatutários não são capturados. A literatura documenta "
 "limitações estruturais de cobertura da CAT (ALMEIDA; BINDER; FISCHER, "
 "2000) e fatores associados à subnotificação como informalidade e "
 "desconhecimento de direitos (GALDINO; SANTANA; FERRITE, 2012).",

 # =======================================================================
 # 3. REFERENCIAL TEÓRICO
 # =======================================================================
 "Oliveira (2004) demonstra que cada regime de acumulação capitalista "
 "produz formas específicas de desgaste da força de trabalho. Antunes e "
 "Praun (2015) mostram que a epidemia de agravos ocupacionais no Brasil é "
 "componente estrutural do padrão de acumulação flexível. No setor saúde, "
 "o trabalho é relacional e corporal: envolve perfurocortantes, fluidos "
 "biológicos, movimentação de pacientes e sofrimento psíquico contínuo. A "
 "OIT estima 2,3 milhões de mortes anuais por acidentes e doenças do "
 "trabalho (ILO, 2023) e a OMS reconhece trabalhadores da saúde como grupo "
 "de alto risco (WHO, 2020)."

 "\n\n"
 "Souza, Melo e Vasconcellos (2017) distinguem o [[i]]campo[[/i]] "
 "institucional da Saúde do Trabalhador - normas, políticas e sistemas de "
 "informação - da [[i]]questão[[/i]] mais ampla dos acidentes que o sistema "
 "não captura. Essa distinção é pertinente a Campos, onde a dualidade "
 "RPPS/INSS produz uma segmentação institucional que determina quais "
 "acidentes se tornam estatisticamente visíveis. Gomez, Vasconcellos e "
 "Machado (2018) apontam que a fragmentação dos sistemas de informação "
 "compromete a capacidade de produzir conhecimento acionável. "

 "\n\n"
 "Este ensaio analítico descreve o perfil e as taxas de notificação de "
 "acidentes de trabalho entre profissionais da saúde com vínculo celetista "
 "em Campos dos Goytacazes, de 2018 a 2025, contextualizando os achados "
 "à luz da literatura sobre sistemas de informação em saúde do trabalhador. "
 "Não se pretende testar hipóteses causais, mas identificar padrões que "
 "contribuam para a discussão sobre vigilância e subnotificação.",

 # =======================================================================
 # 4. MÉTODO
 # =======================================================================
 "A base principal é a CAT do INSS (Portal de Dados Abertos): 58 arquivos "
 "de julho de 2018 a outubro de 2025, dos quais 5.066 vinculados a "
 "empregadores de Campos e 1.144 atribuídos às profissões da saúde celetistas "
 "classificadas pela CBO 2002. O denominador é a RAIS (Ministério do "
 "Trabalho), que registra 83.938 vínculos celetistas ativos em 31 de "
 "dezembro das mesmas categorias, sendo comensurável com a CAT. Ambos "
 "capturam exclusivamente celetistas."

 "\n\n"
 "Bases complementares incluem: SINAN (DATASUS, 72 arquivos, 2018-2025, "
 "filtro SIT_TRAB=01 para acidentes e TRAB_DOE=1 para doenças); SIH/SUS "
 "([[i]]microdatasus[[/i]], R, 255.254 internações com CID de trabalho "
 "baseados na lista do Ministério da Saúde, Brasil, 1999); benefícios "
 "acidentários do INSS, B91-B94 (Portal de Dados Abertos, 48.528 "
 "concessões, todas as ocupações); e 42 indicadores do SmartLab/MPT "
 "(extração automatizada)."

 "\n\n"
 "As taxas de notificação foram calculadas por 1.000 vínculos RAIS, com "
 "intervalos de confiança binomiais exatos a 95%. Células com menos de "
 "cinco eventos ou denominador inferior a 30 foram suprimidas. Três "
 "limitações metodológicas devem ser explicitadas. Primeira, o "
 "delineamento é ecológico: os dados são agregados por categoria "
 "profissional e ano, sem informação individual sobre o regime "
 "previdenciário do trabalhador acidentado (MORGENSTERN, 1995). Segunda, "
 "a cobertura temporal da CAT é irregular (2018 apenas julho-dezembro; "
 "2025 até outubro; lacunas em 2022). Terceira, o SINAN combina arquivos "
 "FINAIS (2018-2022) e PRELIM (2023-2025), cuja completude não é "
 "equivalente. Os resultados referem-se exclusivamente a trabalhadores "
 "celetistas da saúde; não são generalizáveis a estatutários, autônomos "
 "ou pessoas jurídicas.",

 # =======================================================================
 # 5. RESULTADOS
 # =======================================================================
 "Das 1.144 CATs de profissionais celetistas da saúde (Tabela 1), a "
 "enfermagem de nível técnico concentra 70,2% dos registros "
 "(IC 95%: 67,5-72,8%), seguida por enfermeiros (14,2%; IC 95%: "
 "12,2-16,3%). A medicina responde por 1,0% (12 notificações em oito anos; "
 "IC 95%: 0,5-1,7%). Predominam mulheres (85,7%; IC 95%: 83,7-87,7%), "
 "idade mediana de 36 anos e acidentes típicos (81,9%). Ferimentos de "
 "punho e mãos (CID S61, 25,1%) e exposição a doenças transmissíveis "
 "(Z20, 21,9%) lideram os diagnósticos. Agentes infecciosos respondem por "
 "26,9% dos causadores e 95,4% dos empregadores pertencem à CNAE 86-87."

 "\n\n"
 "As taxas de notificação por 1.000 vínculos RAIS (Tabela 2) revelam "
 "heterogeneidade expressiva. A enfermagem técnica apresenta 24,9 "
 "notificações por 1.000 vínculos, seguida por enfermeiros (20,4/1.000) "
 "e técnicos de diagnóstico (16,7/1.000). A medicina registra a menor "
 "taxa entre categorias com denominador suficiente: 1,2 por 1.000 "
 "vínculos. A Figura 1 mostra que o padrão de concentração na enfermagem "
 "técnica se repete em todos os anos, com pico de 64 notificações em 2021, "
 "durante o período mais agudo da pandemia de COVID-19 (VEDOVATO et al., "
 "2021).",

 # =======================================================================
 # 6. DEMAIS SISTEMAS DE INFORMAÇÃO
 # =======================================================================
 "O SINAN registrou 11.634 notificações de agravos relacionados ao "
 "trabalho em Campos (Tabela 3). Os acidentes somaram 10.213 "
 "notificações. As doenças ocupacionais com nexo confirmado somaram "
 "apenas 26 casos em oito anos (9 LER/DORT, 14 transtorno mental, "
 "2 câncer, 1 pneumoconiose), com zero registros de dermatose ocupacional "
 "e PAIR. Galdino, Santana e Ferrite (2012) documentam a subnotificação "
 "expressiva desses agravos no Brasil. A razão entre as 11.634 "
 "notificações do SINAN e as 1.144 CATs da saúde (aproximadamente 10:1) "
 "fornece uma estimativa indireta da magnitude de acidentes que não geram "
 "CAT, embora os universos de cobertura não sejam idênticos.",

 "\n\n"
 "As internações do SIH/SUS (Tabela 4) totalizaram 255.254, das quais "
 "286 (0,11%) com CIDs do escopo da saúde ocupacional: distúrbios "
 "osteomusculares (M75, M65, M79, n=272), dermatoses ocupacionais "
 "(L23-L25, n=11), PAIR (H83.3, n=1) e asma ocupacional (J45.0, n=2). "
 "Foram utilizados exclusivamente códigos com relação direta ao trabalho, "
 "excluindo-se traumatismos, queimaduras, fraturas e intoxicações. "
 "Os distúrbios osteomusculares predominam (95%), compatíveis com a "
 "elevada prevalência de LER/DORT em trabalhadores brasileiros. "
 "Não foram registrados óbitos com esses CIDs no SIM (2018-2024).",

 "\n\n"
 "Os benefícios acidentários do INSS (B91-B94) totalizaram 48.528 "
 "concessões (Tabela 5). O pico de 2020 (10.631) coincide com a pandemia "
 "e pode refletir tanto aumento real de acidentes quanto mudanças "
 "administrativas nos procedimentos de concessão durante a emergência "
 "sanitária. A oscilação entre 286 (2022) e 13.607 (2025) provavelmente "
 "incorpora efeitos de alterações normativas e de capacidade operacional "
 "do INSS, não apenas variação na incidência de acidentes. Estes "
 "benefícios cobrem todas as ocupações do município, não apenas a saúde, "
 "e a base não contém a Classificação Brasileira de Ocupações, impedindo "
 "estratificação por categoria profissional.",

 # =======================================================================
 # 7. DISCUSSÃO
 # =======================================================================
 "Os resultados revelam heterogeneidade substancial nas taxas de "
 "notificação de acidentes de trabalho entre as categorias profissionais "
 "celetistas da saúde em Campos. A taxa de 24,9/1.000 para técnicos de "
 "enfermagem contrasta com 1,2/1.000 para a medicina. Essa diferença "
 "de aproximadamente 21 vezes não pode ser atribuída a um único fator.",

 "\n\n"
 "A literatura oferece múltiplos fatores que podem contribuir, sem que "
 "os dados permitam estabelecer a contribuição relativa de cada um. "
 "O processo de trabalho concentra a execução manual do cuidado na "
 "enfermagem técnica, com maior volume de procedimentos envolvendo "
 "risco perfurocortante e exposição biológica. Diferenças na utilização "
 "de equipamentos de proteção, na exposição temporal a procedimentos de "
 "risco e na propensão a notificar acidentes entre categorias são "
 "documentadas (ALMEIDA; BINDER; FISCHER, 2000; GALDINO; SANTANA; "
 "FERRITE, 2012). Parte dos médicos atua como pessoa jurídica, forma "
 "de vínculo sem CAT. Os estatutários, independentemente da categoria, "
 "estão excluídos do sistema CAT. O delineamento ecológico impede "
 "isolar a contribuição de cada fator.",

 "\n\n"
 "A triangulação com os demais sistemas de informação oferece estimativas "
 "indiretas da subnotificação. A razão SINAN/CAT de aproximadamente 10:1 "
 "é compatível com a existência de um universo de acidentes de trabalho "
 "que não são comunicados via CAT. As 5.585 internações com CID de "
 "ocupacional no SIH, com 286 internações em oito anos (0,11%), "
 "das quais 95% são distúrbios osteomusculares (LER/DORT), reforça "
 "o predomínio desses agravos também na rede hospitalar. A virtual "
 "ausência de doenças ocupacionais crônicas "
 "confirmadas no SINAN (26 casos em oito anos) é consistente com a "
 "dificuldade estrutural de estabelecer nexo entre adoecimento crônico "
 "e trabalho (GALDINO; SANTANA; FERRITE, 2012). Esses padrões "
 "convergentes sugerem que a subnotificação de acidentes e, sobretudo, "
 "de doenças ocupacionais é expressiva em Campos, embora sua magnitude "
 "exata não possa ser quantificada com os dados disponíveis.",

 # =======================================================================
 # 8. LIMITAÇÕES E IMPLICAÇÕES
 # =======================================================================
 "Este ensaio apresenta limitações que devem ser consideradas. O "
 "delineamento ecológico impede inferências sobre mecanismos individuais "
 "(MORGENSTERN, 1995). Os denominadores RAIS restringem-se a celetistas; "
 "os achados não se aplicam a estatutários, autônomos ou PJ. A cobertura "
 "temporal da CAT é irregular. A classificação de CIDs de trabalho no "
 "SIH não valida causalidade ocupacional individual. Os dados do SmartLab "
 "são majoritariamente do Censo 2010. Os benefícios do INSS cobrem "
 "universo distinto do das CATs analisadas. Apesar dessas limitações, a "
 "convergência de múltiplas fontes independentes sugere que os padrões "
 "descritos são robustos.",

 "\n\n"
 "Para a vigilância em saúde do trabalhador, os resultados indicam que "
 "a enfermagem de nível técnico concentra a maior carga de notificações "
 "e apresenta a maior taxa por vínculo. Medidas de proteção dirigidas a "
 "essa categoria - dispositivos de segurança para perfurocortantes e "
 "disponibilidade contínua de EPI - têm potencial de impacto "
 "desproporcional. A integração dos registros do RPPS municipal às bases "
 "nacionais permitiria dimensionar a carga de acidentes entre "
 "estatutários, atualmente invisível ao sistema CAT. Pesquisas futuras "
 "com dados individuais que vinculem o acidente ao regime previdenciário "
 "poderão superar as limitações do delineamento ecológico e quantificar "
 "a contribuição relativa dos múltiplos fatores aqui discutidos.",
]

# ============================ TABELAS ==========================================

T1 = [
 ("Categoria profissional", "n", "%", "IC 95%"),
 ("Enfermagem, técnicos e auxiliares", "803", "70,2", "67,5-72,8"),
 ("Enfermagem, enfermeiros", "163", "14,2", "12,2-16,3"),
 ("Diagnóstico/lab., técnicos", "78", "6,8", "5,4-8,3"),
 ("Fisioterapia", "30", "2,6", "1,7-3,6"),
 ("Farmácia, técnicos", "20", "1,7", "1,0-2,5"),
 ("ACS e afins", "14", "1,2", "0,6-1,9"),
 ("Medicina", "12", "1,0", "0,5-1,7"),
 ("Demais (n<5, agregados)", "24", "2,1", "-"),
 ("Feminino", "980", "85,7", "83,7-87,7"),
]

T2 = [
 ("Categoria", "CATs", "RAIS", "Taxa/1.000"),
 ("Enfermagem, técnicos e auxiliares", "803", "32.263", "24,9"),
 ("Enfermagem, enfermeiros", "163", "7.979", "20,4"),
 ("Diagnóstico/lab., técnicos", "78", "4.664", "16,7"),
 ("Fisioterapia", "30", "3.360", "8,9"),
 ("Farmácia, técnicos", "20", "2.381", "8,4"),
 ("ACS e afins", "14", "3.701", "3,8"),
 ("Medicina", "12", "9.892", "1,2"),
]

T3 = [
 ("Agravo", "Total 2018-2025"),
 ("Acidente de Trabalho", "10.213"),
 ("Acid. c/ Material Biológico", "695"),
 ("Animais Peçonhentos (trabalho)", "700"),
 ("LER/DORT", "9"),
 ("Transtorno Mental", "14"),
 ("Câncer / Pneumoconiose", "2 / 1"),
 ("Dermatose / PAIR", "0 / 0"),
 ("Total", "11.634"),
]

T4 = [
 ("Ano", "Internações", "CID ocupacional"),
 ("2018", "28.941", "22"),
 ("2019", "27.856", "40"),
 ("2020", "25.386", "12"),
 ("2021", "28.579", "6"),
 ("2022", "34.145", "95"),
 ("2023", "35.103", "42"),
 ("2024", "36.998", "40"),
 ("2025", "38.246", "29"),
 ("Total", "255.254", "286"),
]

T5 = [
 ("Ano", "Benefícios (B91-B94)"),
 ("2018", "465"),
 ("2019", "366"),
 ("2020", "10.631"),
 ("2021", "4.061"),
 ("2022", "286"),
 ("2023", "8.660"),
 ("2024", "10.452"),
 ("2025", "13.607"),
 ("Total", "48.528"),
]

# ============================ MONTAGEM =========================================
for i, texto in enumerate(CORPO):
    partes = texto.split("\n\n")
    for parte in partes:
        if parte.strip():
            par(parte, indent=True, after=2)

    if i == 4:
        par("[[b]]Tabela 1.[[/b]] CATs das profissões celetistas da saúde, "
            "Campos dos Goytacazes (RJ), 2018-2025 (n = 1.144).",
            indent=False, size=10, before=4, after=2)
        tabela(T1, "Fonte: CAT/INSS. IC 95% binomial exato. Cobertura: jul/2018-out/2025, parcial em 2018, 2022, 2024-2025.")
        par("[[b]]Tabela 2.[[/b]] Taxa de notificação por 1.000 vínculos RAIS celetistas, "
            "profissões da saúde, Campos (RJ), 2018-2025.",
            indent=False, size=10, before=2, after=2)
        tabela(T2, "Fontes: CAT/INSS e RAIS/MTE. Denominador: vínculos celetistas ativos em 31/12. "
               "Ambos capturam exclusivamente celetistas. Células com n<5 ou denominador<30 suprimidas.")
        figura("saidas/figuras/F1_cat_ano_categorias.png",
               "[[b]]Figura 1.[[/b]] CATs de profissionais celetistas da saúde (n = 1.144) por ano e "
               "categoria. Asterisco = cobertura parcial. Fonte: CAT/INSS.")

    if i == 5:
        par("[[b]]Tabela 3.[[/b]] SINAN, agravos relacionados ao trabalho, Campos (RJ), 2018-2025.",
            indent=False, size=10, before=4, after=2)
        tabela(T3, "Fonte: SINAN/DATASUS (FINAIS 2018-2022 + PRELIM 2023-2025). "
               "Filtro: SIT_TRAB=01 (acidentes), TRAB_DOE=1 (doenças). "
               "Observar que PRELIM têm completude inferior a FINAIS.")
        par("[[b]]Tabela 4.[[/b]] Internações SIH/SUS com CID de trabalho, residentes de Campos (RJ), 2018-2025.",
            indent=False, size=10, before=2, after=2)
        tabela(T4, "Fonte: SIH/SUS (microdatasus, R). CIDs: LER/DORT (M75,M65,M79), dermatoses (L23-L25), "
               "PAIR (H83.3), asma ocupacional (J45.0), exposicao (Z57), fatores trabalho (Y96), "
               "exames ocupacionais (Z10.0). SIM: 0 obitos com esses CIDs.")
        par("[[b]]Tabela 5.[[/b]] Benefícios acidentários (B91-B94) concedidos, Campos (RJ), 2018-2025.",
            indent=False, size=10, before=2, after=2)
        tabela(T5, "Fonte: INSS, Portal de Dados Abertos. Universo: todas as ocupações do município. "
               "Oscilações em 2020-2022 podem refletir mudanças administrativas além de variação epidemiológica. "
               "B91=auxílio-doença; B92=aposentadoria invalidez; B93=pensão morte; B94=auxílio-acidente.")

# ============================ REFERÊNCIAS ======================================
par("[[b]]Referências[[/b]]", indent=False, size=10.5, before=8, after=4)
REFS = [
 "ALMEIDA, I. M.; BINDER, M. C. P.; FISCHER, F. M. Acidentes de trabalho e sua "
 "notificação. [[b]]Revista Brasileira de Saúde Ocupacional[[/b]], v. 25, p. 17-31, 2000.",
 "ANTUNES, R.; PRAUN, L. A sociedade dos adoecimentos no trabalho. [[b]]Serviço Social "
 "& Sociedade[[/b]], n. 123, p. 407-427, 2015.",
 "BRASIL. [[b]]Lei nº 8.213/1991[[/b]]. Planos de Benefícios da Previdência Social.",
 "BRASIL. Ministério da Saúde. [[b]]Portaria nº 1.339/1999[[/b]]. Lista de doenças "
 "relacionadas ao trabalho.",
 "GALDINO, A.; SANTANA, V. S.; FERRITE, S. Fatores associados à subnotificação de "
 "acidentes. [[b]]Cadernos de Saúde Pública[[/b]], v. 28, n. 4, p. 733-742, 2012.",
 "GOMEZ, C. M.; VASCONCELLOS, L. C. F.; MACHADO, J. M. H. Saúde do trabalhador no "
 "SUS. [[b]]Ciência & Saúde Coletiva[[/b]], v. 23, n. 6, p. 1963-1970, 2018.",
 "ILO. [[b]]World Statistic: Occupational Safety and Health[[/b]]. 2023.",
 "LACAZ, F. A. C. O campo Saúde do Trabalhador. [[b]]Cadernos de Saúde Pública[[/b]], "
 "v. 23, n. 4, p. 757-766, 2007.",
 "MARTINS, S.; HASENCLEVER, L.; MIRANDA, C. A gestão da saúde e instabilidade de "
 "financiamento. [[b]]Cadernos do Desenvolvimento Fluminense[[/b]], n. 27, 2024.",
 "MORGENSTERN, H. Ecologic studies in epidemiology. [[b]]Annual Review of Public "
 "Health[[/b]], v. 16, p. 61-81, 1995.",
 "OLIVEIRA, E. M. Transformações no mundo do trabalho. [[b]]Caminhos de Geografia[[/b]], "
 "v. 6, n. 11, p. 84-96, 2004.",
 "SILVA, J. E. M.; HASENCLEVER, L. Ciclo do petróleo e desenvolvimento socioeconômico "
 "em Campos. [[b]]Desenvolvimento em Questão[[/b]], v. 17, n. 46, p. 314-332, 2019.",
 "SOUZA, D. O.; MELO, A. I. S. C.; VASCONCELLOS, L. C. F. Saúde do(s) trabalhador(es): "
 "do 'campo' à 'questão'. [[b]]Saúde em Debate[[/b]], v. 41, n. 113, p. 591-604, 2017.",
 "VEDOVATO, T. G. et al. Trabalhadores(as) da saúde e a COVID-19. [[b]]Revista "
 "Brasileira de Saúde Ocupacional[[/b]], v. 46, e1, 2021.",
 "WHO. [[b]]Occupational health: health workers[[/b]]. 2020.",
]
for r in REFS:
    par(r, indent=False, just=False, size=8, after=1)

# ============================ VERIFICAÇÕES =====================================
conteudo = " ".join(CORPO) + " ".join(REFS)
for l in T1+T2+T3+T4+T5:
    conteudo += " ".join(str(v) for v in l)
for proibido, nome in ((chr(8212), "travessão"), (chr(8211), "meia-risca")):
    if proibido in conteudo:
        raise SystemExit(f"PROIBIDO: {nome} encontrado.")

os.makedirs("documentos", exist_ok=True)
d.save("documentos/artigo.docx")

soffice = shutil.which("soffice") or r"C:\Program Files\LibreOffice\program\soffice.exe"
if os.path.exists(soffice):
    subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir", "documentos",
                    "documentos/artigo.docx"], capture_output=True, timeout=300)
    from pypdf import PdfReader
    npag = len(PdfReader("documentos/artigo.pdf").pages)
    print(f"ensaio.docx gerado com {npag} pagina(s).")
else:
    print("AVISO: LibreOffice ausente.")
