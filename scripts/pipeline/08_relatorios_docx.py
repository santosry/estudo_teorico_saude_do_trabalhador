# -*- coding: utf-8 -*-
"""08_relatorios_docx.py — Relatórios de auditoria (dados, scripts) e metodológico."""
import os
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

RAIZ = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
os.chdir(RAIZ)

def novo_doc():
    d = Document()
    for s in d.sections:
        s.page_width, s.page_height = Cm(21), Cm(29.7)
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2.5)
    st = d.styles["Normal"]
    st.font.name = "Times New Roman"; st.font.size = Pt(11)
    st.paragraph_format.space_after = Pt(4)
    return d

def h(d, t, lvl=1):
    p = d.add_paragraph()
    r = p.add_run(t); r.bold = True
    r.font.size = Pt(13 if lvl == 0 else 11.5)
    p.paragraph_format.space_before = Pt(8)
    return p

def para(d, t, just=True):
    p = d.add_paragraph(t)
    if just: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def tab_from_df(d, df, fsize=9):
    t = d.add_table(rows=1, cols=len(df.columns)); t.style = "Table Grid"
    for j, c in enumerate(df.columns):
        cell = t.rows[0].cells[j]; cell.text = str(c)
        for r_ in cell.paragraphs[0].runs: r_.bold = True; r_.font.size = Pt(fsize)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = "" if pd.isna(v) else str(v)
            for r_ in cells[j].paragraphs[0].runs: r_.font.size = Pt(fsize)
    return t

# ============ RELATÓRIO METODOLÓGICO ============
d = novo_doc()
h(d, "Relatório metodológico — CAT/INSS, Campos dos Goytacazes, profissões da saúde, 2018–2025", 0)
para(d, "Data: 18 jul. 2026. Estudo teórico-conceitual com análise documental e apoio de dados secundários. "
        "Este relatório documenta as decisões metodológicas da reconstrução integral da análise das Comunicações "
        "de Acidente de Trabalho (CAT) do INSS. A etapa SmartLab foi integralmente excluída por decisão metodológica; "
        "nenhum arquivo dessa origem existe na pasta do projeto. A análise legada (medicina x enfermagem) foi tratada "
        "como material catalogado e auditado, sem reaproveitamento de resultados.")
h(d, "1. Fontes")
para(d, "(a) Teóricas: cinco obras da pasta ARTIGOS, lidas integralmente (Antunes, 2018; Cecilio; Lacaz, 2012; Lemos, 2012; "
        "Lourenço, 2012; Oliveira, 2004), complementadas por Mendes e Dias (1991) e Vedovato et al. (2021), verificadas por DOI. "
        "(b) Caracterização municipal: tabelas SIDRA/IBGE da pasta CARACTERIZAÇÃO - CAMPOS (Censo 2022 — tab. 4714; estimativas — 6579; "
        "PIB dos municípios — 5938; PAM — 1612), com auditoria de fonte, período, unidade e preços (correntes). "
        "(c) CAT/INSS: 58 arquivos CSV (competências jul/2018 a dez/2025), 3.902.905 linhas de dados, dicionário oficial da fonte "
        "(dicionario-cat-dados-abertos-10-02-2021.xlsx).")
h(d, "2. Estrutura dos arquivos e importação")
para(d, "Quatro esquemas estruturais foram identificados e mapeados por posição (24 colunas em duas variantes, 25 e 27 colunas), "
        "com cabeçalhos duplicados (CBO, CID-10, CNAE 2.0, Data Acidente) preservados — nenhuma coluna foi renomeada ou "
        "sobrescrita automaticamente. Codificação detectada por arquivo (Latin-1; três arquivos de 2020 em UTF-8 com BOM); separador ponto e vírgula; "
        "sem tratamento de aspas (QUOTE_NONE), como na fonte. A coluna 2 ('Data Acidente') contém o mês de referência do acidente "
        "(formatos AAAA/MM, AAAAMM ou DD/MM/AAAA com dia = 01) e não foi confundida com a data completa do acidente "
        "(posições 21–23, conforme o esquema), usada em todas as análises temporais. Cada linha recebeu identificador técnico "
        "(arquivo:linha) e hash SHA-256 do registro bruto.")
h(d, "3. Critério municipal")
para(d, "Filtro principal: código municipal do EMPREGADOR igual a 330100 (seis dígitos antes do hífen; observado como "
        "'330100-Campos dos Go' truncado e '330100-Campos dos Goytacazes' completo) E UF do município do empregador igual a "
        "Rio de Janeiro. A tabela de controle (dados/processados/controle_localidades_campo.csv) lista 73 localidades com "
        "'Campo(s)' no nome presentes na fonte (Campo Grande-MS, São Bernardo do Campo, São José dos Campos, Campos Novos, "
        "Campos do Jordão, Campos de Júlio, Mário Campos etc.), todas excluídas. Doze registros com código 330100 e UF divergente "
        "foram excluídos e documentados. A unidade observada são CATs vinculadas a empregadores localizados em Campos dos "
        "Goytacazes — não necessariamente acidentes ocorridos no município nem trabalhadores nele residentes.")
h(d, "4. Duplicidades")
para(d, "Arquivos com cobertura sobreposta foram identificados pela distribuição mensal das datas de emissão "
        "(202201/202204; 202207–202211; 202404/202405; 202506–202510). Removeram-se somente duplicidades tecnicamente "
        "demonstráveis: linhas brutas idênticas (mesmo esquema) presentes em mais de um arquivo, mantendo-se a primeira "
        "competência (537 remoções no recorte Campos+RJ; log integral em dados/processados/log_duplicidades_removidas.csv). "
        "Linhas idênticas dentro de um mesmo arquivo foram mantidas e sinalizadas (eventos legítimos podem coincidir).")
h(d, "5. Classificação ocupacional (CBO 2002)")
para(d, "Classificação prioritariamente pelo código de seis dígitos, com dicionário mestre construído a partir das famílias "
        "ocupacionais e revisão manual dos 459 códigos observados na base municipal (metadados/dicionario_cbo_profissoes_saude.xlsx). "
        "Títulos validados contra espelho público da tabela oficial 'CBO2002 – Ocupação' (2.719 títulos; referencias/fonte_cbo.txt). "
        "Três níveis: (i) universo principal (medicina — famílias 2231/2251–2253; enfermagem — 2235/3222; odontologia e saúde bucal; "
        "farmácia; fisioterapia; nutrição; fonoaudiologia; biomedicina; técnicos e auxiliares de diagnóstico e laboratório — 3241/3242/5152; "
        "agentes comunitários e afins — 5151; instrumentação cirúrgica — 322225); (ii) profissões multiprofissionais intersetoriais "
        "(psicologia, serviço social, educação física, biologia), com sensibilidade restrita a CNAE 86/87; (iii) trabalhadores de apoio "
        "em estabelecimentos de saúde (analisados à parte; jamais somados às profissões da saúde). Registros sem CBO válido (184; 3,6%) "
        "foram mantidos em categoria própria. O código 519305 ('enfermeiro veterinário' como sinônimo de auxiliar de veterinário) foi "
        "excluído do campo da saúde humana — erro que a rotina legada, baseada em texto, cometeria.")
h(d, "6. Denominadores e interpretação")
para(d, "As pastas do projeto não continham vínculos formais por ocupação e ano; nenhum indicador de incidência, prevalência, "
        "risco ou taxa foi calculado a partir da CAT isoladamente. Denominadores REAIS de força de trabalho foram baixados do "
        "CNES/DataSUS (TabNet, def cnes/cnv/prid02rj.def): profissionais (indivíduos) por ocupação CBO 2002, município de Campos "
        "(330100), competência dezembro de 2018 a 2025 (9.803 a 13.275 profissionais), com resolução dinâmica e verificada do "
        "filtro municipal, dupla checagem de totais por consulta de controle, brutos preservados em cnes/ e "
        "proveniência em logs/log_10_denominadores.json. Como o CNES abrange TODOS os tipos de vínculo (estatutário, celetista, "
        "autônomo, PJ; SUS e não SUS) e a CAT cobre essencialmente celetistas, a compatibilidade numerador-denominador NÃO é "
        "demonstrável; as razões CAT por 1.000 profissionais CNES (saidas/tabelas/T22_razao_cat_1000_cnes.csv) são apresentadas "
        "exclusivamente como densidade EXPLORATÓRIA de comunicação (supressão quando numerador<5 ou denominador<30), jamais como "
        "incidência ou risco. A RAIS/eSocial (vínculos formais por CBO x município x ano, via PDET/MTE) permanece como denominador "
        "prioritário para trabalho futuro: os microdados exigem download de arquivos de vários GB por UF/ano — bloqueio registrado, "
        "sem qualquer imputação.")
h(d, "7. Cobertura temporal e sensibilidade")
para(d, "A série da fonte inicia na competência jul/2018 (2018 parcial). Em 2022 a carga da fonte é irregular; as competências "
        "set–dez/2024 têm volume nacional atípico (3–10 mil linhas/mês contra 40–60 mil típicas) e nov–dez/2025 estão quase vazias "
        "(205 e 126 linhas): 2025 foi tratado como parcial (efetivo até out). Análises de sensibilidade: exclusão de 2025; meses "
        "jan–out em todos os anos; exclusão de trajeto; somente típicos; universo restrito a CNAE 86/87; multiprofissionais separadas; "
        "resultados antes/depois da deduplicação (saidas/tabelas/T17_sensibilidade.csv).")
h(d, "8. Validação independente e reprodutibilidade")
para(d, "Uma segunda rotina (pandas, implementação distinta) recontou os totais diretamente dos 58 CSVs brutos: totais de linhas, "
        "base municipal, universo principal, categorias e anos convergiram integralmente com o pipeline (logs/validacao_independente.json: "
        "'CONVERGENTE — publicação liberada'). Scripts com caminhos relativos, executáveis em ordem numerada, com logs; versões em "
        "metadados/versoes_programas_pacotes.txt; sem procedimentos aleatórios.")
d.save("documentos/relatorio_metodologico.docx")

# ============ AUDITORIA DE DADOS ============
d = novo_doc()
h(d, "Relatório de auditoria dos dados — CAT/INSS 2018–2025 e caracterização municipal", 0)
para(d, "Síntese dos achados de auditoria (detalhes em logs/log_auditoria.csv e logs/log_qualidade_dados.csv). "
        "Nenhum arquivo SmartLab presente (etapa excluída por decisão metodológica).")
h(d, "1. Achados (origem — problema — gravidade — correção)")
aud = pd.read_csv("logs/log_auditoria.csv", sep=";", encoding="utf-8-sig")
tab_from_df(d, aud[["origem", "problema_evidencia", "gravidade", "correcao_decisao"]], 8)
h(d, "2. Fluxo de seleção de registros")
tab_from_df(d, pd.read_csv("saidas/tabelas/T19_fluxo_selecao.csv", sep=";", encoding="utf-8-sig"), 9)
h(d, "3. Completude (universo principal, % de válidos por ano)")
tab_from_df(d, pd.read_csv("saidas/tabelas/T15_completude_pct.csv", sep=";", encoding="utf-8-sig"), 8)
para(d, "Notas: a queda de completude de CBO em 2021–2023 decorre de '{ñ class}' na fonte; a data de emissão inexiste no esquema "
        "das competências jun–out/2023; agente causador sem informação em cerca de 10% dos registros.")
h(d, "4. Classificação dos resultados")
para(d, "CONFIÁVEIS: fluxo de seleção; distribuição por categoria profissional, sexo, tipo de acidente, parte do corpo, agente, "
        "natureza, CID-10, CNAE e emitente; totais de 2019, 2020, 2021 e 2023 (cobertura integral de competências). "
        "CONFIÁVEIS COM RESSALVAS: comparações anuais envolvendo 2018, 2022, 2024 e 2025 (cobertura parcial ou irregular da fonte); "
        "idade (7 registros sem nascimento); tempo acidente–emissão (ausente em parte de 2023); categorias com n<5 (suprimidas/agregadas). "
        "EXPLORATÓRIOS: comparação de médias mensais entre períodos pré/critico/pós-pandemia; proporção de agente infeccioso por ano; "
        "recorte de trabalhadores de apoio em CNAE saúde; razões CAT/1.000 profissionais CNES (T21/T22 — denominador inclui vínculos "
        "não celetistas; densidade de comunicação, não incidência). NÃO UTILIZÁVEIS: quaisquer resultados da análise legada (universo restrito, "
        "duplicidades não tratadas, classificação textual; saídas antigas não localizadas na pasta — irreprodutíveis).")
d.save("documentos/relatorio_auditoria_dados.docx")

# ============ AUDITORIA DE SCRIPTS ============
d = novo_doc()
h(d, "Relatório de auditoria dos scripts legados", 0)
para(d, "Scripts auditados: script legado da CAT (R; data.table/dplyr/ggplot2; pipeline Enfermagem x Medicina; original 'CAT - INSS/script.R') e "
        "script legado da caracterização (R; sidrar; original 'CARACTERIZAÇÃO - CAMPOS/script.R'). Originais preservados em scripts/legado/; cópias comentadas "
        "no mesmo diretório (marcadores [A1]–[A8], [B1]–[B3]); alterações em logs/log_alteracoes_scripts.csv. O ambiente R não está "
        "disponível nesta máquina (Rscript ausente) e a pasta RESULTADOS/ das execuções antigas não existe: os resultados legados são "
        "irreprodutíveis e foram descartados como fonte. Nenhuma menção a SmartLab foi encontrada nos scripts.")
h(d, "Achados por script")
alt = pd.read_csv("logs/log_alteracoes_scripts.csv", sep=";", encoding="utf-8-sig")
tab_from_df(d, alt[["script", "linha_bloco", "problema", "gravidade", "resultado_anterior", "resultado_corrigido"]], 8)
h(d, "Verificações específicas solicitadas")
para(d, "Filtro textual 'Campos' sem validação de código: NÃO ocorre na forma mais grave (o legado usa código OU nome completo "
        "'campos dos goytacazes'), mas não valida UF do empregador (12 registros afetados). UF do acidente em lugar da UF do empregador: "
        "não ocorre. Confusão competência x data completa: mitigada por regex de formato, porém frágil ante mudanças de esquema. "
        "Cabeçalhos duplicados sobrescritos: renomeados automaticamente (make.unique), com risco de mistura entre esquemas. "
        "Perda de zeros em CBO: não ocorre (leitura como texto). Filtro antigo restrito a medicina/enfermagem: confirmado (perda de 166 "
        "registros de outras categorias). Perda de técnicos/auxiliares: não ocorre dentro da enfermagem (família 3222 incluída). "
        "Contagem duplicada por sobreposição de arquivos: confirmada (537 duplicidades no recorte municipal; corrigida). Exclusão "
        "silenciosa de não classificados: confirmada (filtro final descarta CBO ausente sem contabilizar; corrigida). Mistura de anos "
        "completos e incompletos: confirmada (sem tratamento de 2018/2022/2024/2025; corrigida). Frequências como risco: o próprio "
        "legado adverte contra, mas aplica testes qui-quadrado sem denominadores (não replicados).")
d.save("documentos/relatorio_auditoria_scripts.docx")
print("3 relatórios DOCX gerados em documentos/")
